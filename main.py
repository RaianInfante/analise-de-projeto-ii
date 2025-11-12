import discord
import re
import time
import os
import aiohttp
import asyncio
import tempfile
from discord.ui import View, Button
from faster_whisper import WhisperModel
from openai import OpenAI
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
from dotenv import load_dotenv
from datetime import datetime


load_dotenv()

# Carregar token e chave API
DISCORD_TOKEN = os.getenv("DISCORD_TOKEN")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
client_ia = OpenAI(api_key=OPENAI_API_KEY)

# Configurações do Discord
intents = discord.Intents.default()
intents.message_content = True
client = discord.Client(intents=intents)

# Modelo Whisper local
print("🔄 Carregando modelo Whisper (pode demorar)...")
model = WhisperModel("small")
print("✅ Modelo carregado")


async def summarize_text(text: str) -> str:
    """Gera um resumo da reunião usando a nova API da OpenAI."""
    prompt = (
        "Você é um assistente que resume reuniões. "
        "Dado o texto da reunião abaixo, gere:\n"
        "1. Um resumo conciso da reunião.\n"
        "2. Principais tópicos discutidos (em lista).\n"
        "3. Principais decisões ou ações a tomar (em lista).\n\n"
        f"Texto da reunião:\n{text}\n\n"
        "Formato da resposta:\n"
        "Resumo:\n<texto>\n\nTópicos:\n- Tópico1\n- Tópico2\n\nAções/Decisões:\n- Ação1\n- Ação2"
    )

    try:
        response = client_ia.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system",
                    "content": "Você é um assistente especialista em resumir reuniões."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=500
        )

        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"Erro ao gerar resumo: {e}"


@client.event
async def on_ready():
    print(f"🤖 Logado como {client.user}")

    canal_id = 1414730905720979620
    channel = client.get_channel(canal_id)

    if channel:
        await channel.send("✅ VoxNote Conectado e pronto para uso!")


class CanalButtonEntrar(View):
    def __init__(self, canais, autor):
        super().__init__(timeout=30)
        self.autor = autor

        # Cria um botão para cada canal de voz
        for canal in canais:
            botao = Button(label=canal.name, style=discord.ButtonStyle.primary)
            botao.callback = self.make_callback(canal)
            self.add_item(botao)

        # Botão de cancelar
        botao_cancelar = Button(
            label="❌ Cancelar", style=discord.ButtonStyle.danger)
        botao_cancelar.callback = self.cancelar
        self.add_item(botao_cancelar)

    def make_callback(self, canal):
        async def callback(interaction: discord.Interaction):
            # Impede outros usuários de usar o menu
            if interaction.user != self.autor:
                return

            # Conecta o bot no canal selecionado
            await canal.connect()

            # Move o usuário junto, se possível
            if interaction.user.voice:
                try:
                    await interaction.user.move_to(canal)
                    await interaction.response.edit_message(
                        content=f"🎤 Conectei e movi você para o canal: **{canal.name}** ✅",
                        view=None
                    )
                except discord.Forbidden:
                    await interaction.response.edit_message(
                        content=f"🎤 Entrei no canal **{canal.name}**, mas não tenho permissão para mover você.",
                        view=None
                    )
            else:
                await interaction.response.edit_message(
                    content=f"🎤 Conectei no canal **{canal.name}**!",
                    view=None
                )
        return callback

    async def cancelar(self, interaction: discord.Interaction):
        """Função do botão Cancelar"""
        if interaction.user != self.autor:
            return

        await interaction.response.edit_message(
            content="❌ Ação cancelada.",
            view=None
        )


class CanalButtonTrocar(View):
    def __init__(self, autor):
        super().__init__(timeout=30)
        self.autor = autor

        botaoSim = Button(label="✅ Sim", style=discord.ButtonStyle.success)
        botaoSim.callback = self.sim
        self.add_item(botaoSim)

        botaoNao = Button(label="❌ Não", style=discord.ButtonStyle.danger)
        botaoNao.callback = self.nao
        self.add_item(botaoNao)

    async def sim(self, interaction: discord.Interaction):
        """Se o usuário clicar em 'Sim', o bot troca para o canal do usuário"""
        if interaction.user != self.autor:
            return

        # Verifica se o usuário está em um canal de voz
        if not interaction.user.voice:
            await interaction.response.edit_message(
                content="⚠️ Você não está em nenhum canal de voz para eu me mover.",
                view=None
            )
            return

        canal_usuario = interaction.user.voice.channel

        # Desconecta o bot atual e reconecta no novo canal
        voice_client = discord.utils.get(
            interaction.client.voice_clients, guild=interaction.guild)
        if voice_client:
            await voice_client.disconnect()

        await canal_usuario.connect()
        await interaction.response.edit_message(
            content=f"🔄 Conectei ao seu canal de voz: **{canal_usuario.name}** ✅",
            view=None
        )

    async def nao(self, interaction: discord.Interaction):
        """Se o usuário clicar em 'Não', o menu fecha"""
        if interaction.user != self.autor:
            return

        await interaction.response.edit_message(
            content="❌ Ação cancelada.",
            view=None
        )


@client.event
async def on_message(message: discord.Message):

    global processando

    if message.author == client.user:
        return

    comando = message.content.lower()

    # 📜 Exibe lista de comandos se o usuário digitar "!comandos"
    if comando.startswith("!comandos"):
        comandos = (
            "**🤖 Comandos disponíveis:**\n\n"
            "📥 `!ajuda` — Mostra como enviar um áudio para transcrição.\n"
            "🎧 `!entrar` — Faz o bot entrar em um canal de voz (ou mostra os disponíveis).\n"
            "🔊 `!sair` — Faz o bot sair do canal de voz atual.\n"
            "📤 `!desconectar` — Desliga completamente o bot (somente admin).\n"
            "\n💡 *Envie um áudio (.mp3, .wav, .m4a...) e eu transcrevo e gero um resumo completo!*"
        )

        embed = discord.Embed(
            title="📚 Central de Comandos — VoxNote",
            description=comandos,
            color=discord.Color.blue()
        )
        embed.set_footer(
            text="Digite o comando desejado, por exemplo: !entrar")

        await message.channel.send(embed=embed)
        return

    if comando.startswith("!entrar"):
        # Verifica se o autor está em algum canal de voz
        if message.author.voice:
            canal = message.author.voice.channel

            # Se já estiver conectado, evita duplicação
            if discord.utils.get(client.voice_clients, guild=message.guild):
                view = CanalButtonTrocar(message.author)
                await message.channel.send(
                    "⚠️ Já estou conectado a um canal de voz neste servidor! Deseja que eu troque para o seu?",
                    view=view
                )
                return

            await canal.connect()
            await message.channel.send(f"🎧 Conectado ao canal: **{canal.name}** ✅")

        else:
            # Se não estiver em canal, mostra botões com canais disponíveis
            guild = message.guild
            canais_voz = [c for c in guild.voice_channels]

            if not canais_voz:
                await message.channel.send("⚠️ Não há canais de voz disponíveis neste servidor.")
                return

            view = CanalButtonEntrar(canais_voz, message.author)
            await message.channel.send(
                "🎵 Você não está em um canal de voz.\nEscolha abaixo o canal que deseja que eu entre:",
                view=view
            )

    if comando.startswith("!sair"):
        if client.voice_clients:
            voice_client = client.voice_clients[0]
            canal_nome = voice_client.channel.name

            await voice_client.disconnect()
            await message.channel.send(f"👋 Desconectado do canal de voz: **{canal_nome}**")
        else:
            await message.channel.send("⚠️ Eu não estou conectado a nenhum canal de voz no momento.")

    # Comando de ajuda
    if comando.startswith("!ajuda"):
        await message.channel.send(
            "🎙️ Envie um arquivo de áudio (.mp3, .wav, .m4a, .flac...) que eu transcrevo e gero o resumo da reunião!"
        )
        return

    if comando.startswith("!desconectar"):

        dono_id = 756982258489098332

        if message.author.id != dono_id:
            await message.channel.send("🚫 Você não tem permissão para me desconectar.")
            return

        await message.channel.send("👋 Desconectando... Até logo!")
        await client.close()

    # Verifica anexos
    if message.attachments:
        for attachment in message.attachments:
            if attachment.filename.lower().endswith((".mp3", ".wav", ".m4a", ".flac", ".ogg", ".aac", ".aup", "flac")):
                await message.channel.send("🎧 Recebi seu áudio! Transcrevendo... ⏳")

                # Download do áudio

                inicio = time.time()

                file_path = os.path.join(
                    tempfile.gettempdir(), attachment.filename)
                async with aiohttp.ClientSession() as session:
                    async with session.get(attachment.url) as resp:
                        with open(file_path, "wb") as f:
                            f.write(await resp.read())

                # Transcrição
                try:
                    segments, info = model.transcribe(file_path, beam_size=5)
                    full_text = "".join(segment.text for segment in segments)
                except Exception as e:
                    await message.channel.send(f"❌ Erro ao transcrever: {e}")
                    return

                # Salva a transcrição
                txt_path = file_path.rsplit(".", 1)[0] + ".txt"
                with open(txt_path, "w", encoding="utf-8") as txt_file:
                    txt_file.write(full_text)

                await message.channel.send("✅ Transcrição concluída. Gerando resumo... ⏳")

                # Gera resumo
                summary = await summarize_text(full_text)

                fim = time.time()
                duracao = fim - inicio
                segundos = int(duracao % 60)

                # Formata o resumo em um arquivo TXT bonito
                formatted_summary = (
                    "Resumo da Reunião\n"
                    "=================\n\n"
                    f"{summary}\n"
                )

                # Cria documento Word
                doc = Document()
                doc.add_heading("Resumo da Reunião", level=1)

                # Adiciona conteúdo formatado
                doc.add_paragraph(
                    "Resumo:", style='List Bullet').runs[0].bold = True
                doc.add_paragraph(summary.split("Tópicos:")[
                                  0].replace("Resumo:", "").strip())

                # Adiciona Tópicos
                if "Tópicos:" in summary:
                    doc.add_heading("Tópicos Principais", level=2)
                    topics_part = summary.split(
                        "Tópicos:")[1].split("Ações/Decisões:")[0]
                    for line in topics_part.split("\n"):
                        if line.strip().startswith("-"):
                            doc.add_paragraph(
                                line.strip(), style='List Bullet')

                # Adiciona Ações / Decisões
                if "Ações/Decisões:" in summary:
                    doc.add_heading("Ações / Decisões", level=2)
                    actions_part = summary.split("Ações/Decisões:")[1]
                    for line in actions_part.split("\n"):
                        if line.strip().startswith("-"):
                            doc.add_paragraph(
                                line.strip(), style='List Bullet')

                # Diretório base do arquivo original
                base_dir = os.path.dirname(file_path)

                agora = datetime.now().strftime("%Y%m%d_%H.%M.%S")

                usuario = re.sub(r"[^a-zA-Z0-9_-]+", "_",
                                 message.author.name).strip("_")

                nome_base = f"reuniao_{usuario}_{agora}"

                summary_path = os.path.join(base_dir, f"{nome_base}.docx")
                pdf_path = os.path.join(base_dir, f"{nome_base}.pdf")

                # Salva documento e converte para PDF
                doc.save(summary_path)
                convert(summary_path, pdf_path)

                fim = time.time()
                duracao = fim - inicio
                segundos = int(duracao % 60)

                await message.channel.send(f"📄 Aqui está o resumo da reunião - {segundos}s (formato PDF):", file=discord.File(pdf_path))

                # Também mostra o texto do resumo direto no chat (opcional)
                await message.channel.send(f"🗒️ **Resumo (versão curta):**\n{summary[:1900]}")

                # Limpeza
                os.remove(file_path)
                os.remove(summary_path)

                return

client.run(DISCORD_TOKEN)
